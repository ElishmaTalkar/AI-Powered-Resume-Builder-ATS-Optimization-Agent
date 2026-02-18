import os
import subprocess
import jinja2
from docx import Document
import shutil

class PDFGenerator:
    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")
    OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")

    def __init__(self):
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(self.TEMPLATE_DIR),
            block_start_string='\BLOCK{', # Avoid conflict with latex {}
            block_end_string='}',
            variable_start_string='\VAR{',
            variable_end_string='}',
            comment_start_string='\#{',
            comment_end_string='}',
            line_statement_prefix='%%',
            line_comment_prefix='%#',
            trim_blocks=True,
            autoescape=False,
        )


    def generate_resume(self, data: dict, format: str = "pdf", template_name: str = "classic") -> str:
        if format == "pdf":
            return self._generate_pdf(data, template_name)
        elif format == "docx":
            return self._generate_docx(data)
        else:
            raise ValueError("Unsupported format")

    def escape_latex(self, text: str) -> str:
        """Escape LaTeX special characters."""
        if not isinstance(text, str):
            return text
        chars = {
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
            "\\": r"\textbackslash{}",
        }
        escaped = ""
        for char in text:
            escaped += chars.get(char, char)
        return escaped

    def sanitize_data(self, data):
        """Recursively escape strings in data for LaTeX."""
        if isinstance(data, dict):
            return {k: self.sanitize_data(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self.sanitize_data(v) for v in data]
        elif isinstance(data, str):
            return self.escape_latex(data)
        else:
            return data

    def _generate_pdf(self, data: dict, template_name: str) -> str:
        try:
            # save raw name for filename
            raw_name = data.get('name', 'user')
            
            # Escape special LaTeX characters in data (returns new dict)
            sanitized_data = self.sanitize_data(data)

            # Ensure template exists, default to classic if not found
            if not os.path.exists(os.path.join(self.TEMPLATE_DIR, f"{template_name}.tex")):
                template_name = "classic"
            
            template = self.jinja_env.get_template(f"{template_name}.tex")
            rendered_tex = template.render(**sanitized_data)
            
            # Sanitize filename
            safe_name = "".join([c if c.isalnum() else "_" for c in raw_name])
            tex_filename = f"resume_{safe_name}.tex"
            tex_path = os.path.join(self.OUTPUT_DIR, tex_filename)
            
            with open(tex_path, "w") as f:
                f.write(rendered_tex)
            
            # Compile with pdflatex
            # Check if pdflatex is available in PATH, if not check common Mac locations
            if not shutil.which("pdflatex"):
                # Common MacTeX paths
                possible_paths = ["/Library/TeX/texbin", "/usr/local/bin", "/usr/texbin"]
                for p in possible_paths:
                    if os.path.exists(os.path.join(p, "pdflatex")):
                        os.environ["PATH"] += os.pathsep + p
                        break
            
            if not shutil.which("pdflatex"):
                 raise EnvironmentError("pdflatex not found. Please install TeX distribution.")

            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-halt-on-error", "-output-directory", self.OUTPUT_DIR, tex_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            
            pdf_filename = tex_filename.replace(".tex", ".pdf")
            return os.path.join(self.OUTPUT_DIR, pdf_filename)
        except Exception as e:
             raise RuntimeError(f"PDF Generation failed: {str(e)}")

    def _get_minimal_template(self):
        return r"""
\documentclass[11pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[empty]{fullpage}
\usepackage[hidelinks]{hyperref}
\usepackage{titlesec}

% Minimal formatting
\titleformat{\section}{\large\bfseries\uppercase}{}{0em}{}[\titlerule]

\begin{document}

\begin{center}
    {\huge \textbf{VAR_NAME}} \\
    VAR_EMAIL | VAR_PHONE | VAR_LOCATION
\end{center}

\section{Summary}
VAR_SUMMARY

\section{Experience}
VAR_EXPERIENCE

\section{Education}
VAR_EDUCATION

\section{Skills}
VAR_SKILLS

\end{document}
"""

    def _generate_docx(self, data: dict) -> str:
        doc = Document()
        doc.add_heading(data.get('name', 'Name'), 0)
        
        doc.add_paragraph(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}")
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(data.get('summary', ''))
        
        doc.add_heading('Experience', level=1)
        for job in data.get('experience', []):
            p = doc.add_paragraph()
            p.add_run(f"{job.get('role')} at {job.get('company')}").bold = True
            p.add_run(f"\n{job.get('dates')} | {job.get('location')}")
            for detail in job.get('details', []):
                doc.add_paragraph(detail, style='List Bullet')
                
        doc.add_heading('Education', level=1)
        for edu in data.get('education', []):
            doc.add_paragraph(f"{edu.get('degree')} - {edu.get('school')}")
            doc.add_paragraph(f"{edu.get('dates')} | {edu.get('location')}")

        filename = f"resume_{data.get('name', 'user').replace(' ', '_')}.docx"
        file_path = os.path.join(self.OUTPUT_DIR, filename)
        doc.save(file_path)
        return file_path
